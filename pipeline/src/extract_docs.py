import os
import sys
import json
import asyncio
import re
from pathlib import Path
import docx
from pypdf import PdfReader
from google import genai
from google.genai import types

# Ensure we can import from src/
sys.path.insert(0, str(Path(__file__).resolve().parent))
from schema_types import ExtractedDocument

MEETINGS_DIR = Path(__file__).resolve().parents[2] / "original-docs"
OUTPUT_DIR = Path(__file__).resolve().parents[1] / "data" / "extracted_docs"

# Initialize Gen AI Client
client = genai.Client()
MODEL_ID = "gemini-2.5-flash"

EXTRACTION_PROMPT = """You are an information extraction assistant. Extract structured metadata from this document.

Return valid JSON matching the requested schema.

Rules:
- "docId": a unique kebab-case identifier for the document, e.g. "doc-9-faqs" or "doc-lindy-research".
- "title": the title of the document.
- "date": the date of creation, publication, or the most relevant date of the document in YYYY-MM-DD format. If no date is found, leave it null.
- "authors": list of authors, creators, or main individuals/parties responsible for or mentioned in the document.
- "topics": provide a list of relevant topics or tags (e.g. kebab-case or simple labels like marketing, website-refresh, outbound-sales).
- "summary": a brief summary of the document (1-2 sentences).
"""

def extract_text_from_docx(file_path: Path) -> str:
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text.append(para.text)
    return '\n'.join(full_text)

def extract_text_from_pdf(file_path: Path) -> str:
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text_parts.append(t)
    return "\n".join(text_parts)

def extract_text_from_txt(file_path: Path) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

async def extract_document_metadata(text: str, filename: str) -> ExtractedDocument:
    prompt = f"{EXTRACTION_PROMPT}\n\nFilename: '{filename}'\n\nDocument text:\n\n{text}"

    # Use run_in_executor to avoid blocking the event loop for sync API calls
    loop = asyncio.get_event_loop()
    response = await loop.run_in_executor(
        None,
        lambda: client.models.generate_content(
            model=MODEL_ID,
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=ExtractedDocument,
                temperature=0.1
            )
        )
    )
    
    try:
        data = ExtractedDocument.model_validate_json(response.text)
        # Make sure the text field is populated (Pydantic validates it, but we set it in the process step)
        return data
    except Exception as e:
        print(f"Error parsing validation schema for {filename}: {e}")
        cleaned = response.text.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.replace("```json", "").replace("```", "").strip()
        parsed = json.loads(cleaned)
        # Add text field temporarily to satisfy validation if missing
        if "text" not in parsed:
            parsed["text"] = text
        return ExtractedDocument(**parsed)

async def process_file(filename: str, force: bool = False) -> bool:
    file_path = MEETINGS_DIR / filename
    output_id = Path(filename).stem
    output_path = OUTPUT_DIR / f"{output_id}.json"

    if output_path.exists() and not force:
        print(f"  Skipping (already extracted): {filename}")
        return True

    print(f"  Extracting: {filename}")
    try:
        ext = file_path.suffix.lower()
        if ext == ".md" or ext == ".txt":
            text = extract_text_from_txt(file_path)
        elif ext == ".pdf":
            text = extract_text_from_pdf(file_path)
        elif ext == ".docx":
            text = extract_text_from_docx(file_path)
        else:
            print(f"  Unsupported file extension: {ext}")
            return False

        # Truncate very long texts to respect rate limits / token limits
        original_text = text
        if len(text) > 40000:
            print(f"    Truncating {filename} from {len(text)} to 40000 chars for metadata call")
            text = text[:40000] + "\n\n[TRUNCATED]"

        data = await extract_document_metadata(text, filename)
        # Always set the full text to the actual extracted text, not the truncated one
        data.text = original_text
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(data.model_dump_json(indent=2))
        return True
    except Exception as e:
        print(f"  FAILED to process {filename}: {e}")
        import traceback
        traceback.print_exc()
        return False

async def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if not MEETINGS_DIR.exists():
        print(f"Error: Original documents directory does not exist: {MEETINGS_DIR}")
        return

    all_files = os.listdir(MEETINGS_DIR)
    valid_extensions = {".pdf", ".docx", ".txt", ".md"}
    
    # Filter out meetings (which contain 'notes by gemini' or 'leadership call') and readmes
    candidate_files = []
    for f in all_files:
        ext = Path(f).suffix.lower()
        is_meeting = "notes by gemini" in f.lower() or "leadership call" in f.lower()
        if ext in valid_extensions and not is_meeting and f.lower() != "readme.md":
            candidate_files.append(f)

    # Deduplicate: if both DOCX and PDF exist for a document, only keep the DOCX
    groups = {}
    for f in candidate_files:
        stem = Path(f).stem
        if stem.lower().endswith(".docx"):
            stem = stem[:-5]
        norm_stem = re.sub(r'\s+', ' ', stem).strip().lower()
        groups.setdefault(norm_stem, []).append(f)
        
    deduped_files = []
    for norm_stem, files in groups.items():
        best_file = None
        for ext in [".docx", ".md", ".txt", ".pdf"]:
            match = [f for f in files if f.lower().endswith(ext)]
            if match:
                best_file = match[0]
                break
        if best_file:
            deduped_files.append(best_file)

    print(f"Found {len(deduped_files)} unique non-meeting documents to process.\n")

    force = "--force" in sys.argv
    
    # Filter files that need to be processed
    to_process = []
    for f in deduped_files:
        output_path = OUTPUT_DIR / f"{Path(f).stem}.json"
        if not output_path.exists() or force:
            to_process.append(f)

    if not to_process:
        print("All files already extracted. Use --force to re-extract.")
        return

    print(f"Processing {len(to_process)} files...\n")

    # Process in batches to respect rate limits
    BATCH_SIZE = 3
    processed = 0
    failed = 0

    for i in range(0, len(to_process), BATCH_SIZE):
        batch = to_process[i:i+BATCH_SIZE]
        tasks = [process_file(f, force=force) for f in batch]
        results = await asyncio.gather(*tasks)
        
        for success in results:
            if success:
                processed += 1
            else:
                failed += 1
        
        print(f"  Progress: {processed + failed}/{len(to_process)} ({failed} failed)\n")
        # Small delay between batches to respect rate limits
        await asyncio.sleep(1.0)

    print(f"\nExtraction complete: {processed} succeeded, {failed} failed.")
    print(f"Output directory: {OUTPUT_DIR}")

if __name__ == "__main__":
    asyncio.run(main())
