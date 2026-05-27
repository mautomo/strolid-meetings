import os
import sys
import json
import asyncio
import re
from pathlib import Path
import docx
from pypdf import PdfReader
from google import genai
from google.genai import types
from pydantic import BaseModel

# Ensure we can import from src/
sys.path.insert(0, str(Path(__file__).resolve().parent))
from schema_types import ExtractedMeeting

MEETINGS_DIR = Path(__file__).resolve().parents[2] / "original-docs"
OUTPUT_DIR = Path(__file__).resolve().parents[1] / "data" / "extracted"

# Initialize Gen AI Client
# Pick up GEMINI_API_KEY from environment
client = genai.Client()
MODEL_ID = "gemini-2.5-flash"

EXCLUDE_FILES = {
    "9 FAQs_ Why Dealers Should Consider Strolid.docx",
    "Jake and Link Marketing Hooks.docx",
    "Marketing hooks 2 more about people.docx",
    "Marketing hooks 3.docx",
    "New marketing hooks.docx",
    "Sophia Outbound and Service inbound Plan.docx",
    "Strolid Website.docx",
    "Website Additions, subtractions and changes.docx",
    "README.md",
}

EXTRACTION_PROMPT = """You are a meeting intelligence analyst. Extract structured data from this meeting transcript/notes.

Return valid JSON matching the requested schema.

Rules:
- Extract decisions, action items, and contributions. Keep all descriptions concise (under 15 words) to stay within output limits.
- Limit output to at most 12 decisions, 12 action items, and 15 contributions. Choose the most important/strategic ones.
- For attendees, use full names. If only first names are used, include what's available.
- "confidence" should be "firm" if explicitly agreed, "tentative" if conditionally agreed, "exploratory" if just discussed.
- For action items, always try to identify an owner. If no owner is clear, use "Unassigned".
- Topics should be consistent kebab-case labels that could be reused across meetings (e.g., "website-refresh", "messaging-strategy", "buyer-targeting").
- Tensions include any pushback, disagreement, concern, or unresolved debate.
- referencesToPast includes any mention of prior meetings, earlier decisions, or "we discussed this before" type references.
- If the meeting notes include a transcript section, extract from BOTH the summary/details AND the transcript.
- "contributions": Identify specific contributions (ideas, topics, or concepts) proposed or presented by individuals. For each:
  - "person_name": full name of the person (use resolved full names like Vinnie Micciche, Michael Donovan, Joe Furnari, Paulo Trovao, Jason Branham, Matt Watson, Sergey, Sophia, Jake, Link).
  - "topic": consistent kebab-case topic label.
  - "description": summary of the contribution.
  - "type": "topic", "idea", or "concept".
  - "occurrence": determine if this contribution/topic/idea is being presented for the "first time", or "repeated" in this meeting.
  - "status": the status of this contribution in the meeting: "approved" (explicitly accepted), "denied" (explicitly rejected), "completed-success" (marked as done/successful), "proposed" (just suggested), or "pending" (left unresolved).
"""

def classify_meeting_type(filename: str) -> str:
    lower = filename.lower()
    if "leadership" in lower:
        return "leadership"
    if "marketing" in lower or "vconic" in lower:
        return "marketing"
    if "product" in lower or "roadmap" in lower:
        return "product"
    if any(k in lower for k in [
        "1-on-1", "1 on 1", "joe and michael", "michael _",
        "_ michael", "vin and michael", "call with michael", "michael and vin"
    ]):
        return "one-on-one"
    if "standup" in lower or "scrum" in lower:
        return "standup"
    if any(k in lower for k in ["planning", "prioritization", "alignment"]):
        return "strategy"
    return "other"

def extract_text_from_docx(file_path: Path) -> str:
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text.append(para.text)
    return '\n'.join(full_text)

def extract_text_from_pdf(file_path: Path) -> str:
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text_parts.append(t)
    return "\n".join(text_parts)

async def extract_meeting_data(text: str, filename: str) -> ExtractedMeeting:
    meeting_type = classify_meeting_type(filename)
    prompt = f"{EXTRACTION_PROMPT}\n\nHint: This appears to be a '{meeting_type}' type meeting based on the filename: '{filename}'\n\nMeeting notes:\n\n{text}"

    try:
        response = await asyncio.to_thread(
            client.models.generate_content,
            model=MODEL_ID,
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=ExtractedMeeting,
                temperature=0.1,
                max_output_tokens=8192
            )
        )
        try:
            data = ExtractedMeeting.model_validate_json(response.text)
            return data
        except Exception as e:
            # Fallback to loading text as dict then validating
            cleaned = response.text.strip()
            if cleaned.startswith("```"):
                cleaned = cleaned.replace("```json", "").replace("```", "").strip()
            parsed = json.loads(cleaned)
            return ExtractedMeeting(**parsed)
    except Exception as first_error:
        print(f"  [RETRYING] First extraction attempt failed/truncated for {filename}: {first_error}. Retrying with strict brevity constraints...")
        retry_prompt = (
            f"{EXTRACTION_PROMPT}\n\n"
            "CRITICAL: The previous attempt resulted in a truncated response because it was too long. "
            "To prevent truncation, you MUST restrict the output to a strict maximum of 5 decisions, 5 action items, "
            "and 5 contributions. Keep all description values extremely short (maximum 10 words each).\n\n"
            f"Hint: This appears to be a '{meeting_type}' type meeting based on the filename: '{filename}'\n\n"
            f"Meeting notes:\n\n{text}"
        )
        
        response = await asyncio.to_thread(
            client.models.generate_content,
            model=MODEL_ID,
            contents=retry_prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=ExtractedMeeting,
                temperature=0.1,
                max_output_tokens=8192
            )
        )
        try:
            data = ExtractedMeeting.model_validate_json(response.text)
            return data
        except Exception as e:
            cleaned = response.text.strip()
            if cleaned.startswith("```"):
                cleaned = cleaned.replace("```json", "").replace("```", "").strip()
            parsed = json.loads(cleaned)
            return ExtractedMeeting(**parsed)

async def process_file(filename: str, force: bool = False) -> bool:
    file_path = MEETINGS_DIR / filename
    output_id = Path(filename).stem
    output_path = OUTPUT_DIR / f"{output_id}.json"

    if output_path.exists() and not force:
        print(f"  Skipping (already extracted): {filename}")
        return True

    print(f"  Extracting: {filename}")
    try:
        if filename.endswith(".md"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        elif filename.endswith(".pdf"):
            text = extract_text_from_pdf(file_path)
        else:
            text = extract_text_from_docx(file_path)

        # Truncate very long transcripts to respect limits
        if len(text) > 50000:
            print(f"    Truncating {filename} from {len(text)} to 50000 chars")
            text = text[:50000] + "\n\n[TRUNCATED - remaining transcript omitted]"

        data = await extract_meeting_data(text, filename)
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(data.model_dump_json(indent=2))
        print(f"  Successfully extracted: {filename}")
        return True
    except Exception as e:
        print(f"  FAILED to process {filename}: {e}")
        import traceback
        traceback.print_exc()
        return False

async def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if not MEETINGS_DIR.exists():
        print(f"Error: Original documents directory does not exist: {MEETINGS_DIR}")
        return

    all_files = os.listdir(MEETINGS_DIR)
    meeting_files = [
        f for f in all_files
        if (f.endswith(".docx") or f.endswith(".pdf") or f.endswith(".md"))
        and ("notes by gemini" in f.lower() or "leadership call" in f.lower())
        and f not in EXCLUDE_FILES
    ]

    # Deduplicate: if both DOCX/MD and PDF exist, only keep DOCX/MD
    groups = {}
    for f in meeting_files:
        stem = Path(f).stem
        # Clean any trailing copy numbers like " (1)" or " (2)"
        stem = re.sub(r'\s*\(\d+\)$', '', stem).strip()
        groups.setdefault(stem, []).append(f)
        
    deduped_meetings = []
    for stem, files in groups.items():
        best_file = None
        for ext in [".docx", ".md", ".pdf"]:
            match = [f for f in files if f.lower().endswith(ext)]
            if match:
                best_file = match[0]
                break
        if best_file:
            deduped_meetings.append(best_file)
            
    meeting_files = deduped_meetings
    print(f"Found {len(meeting_files)} unique meeting files to process.\n")

    force = "--force" in sys.argv
    
    # Filter files that need to be processed
    to_process = []
    for f in meeting_files:
        output_path = OUTPUT_DIR / f"{Path(f).stem}.json"
        if not output_path.exists() or force:
            to_process.append(f)

    if not to_process:
        print("All files already extracted. Use --force to re-extract.")
        return

    print(f"Processing {len(to_process)} files...\n")

    # Process in batches to respect rate limits
    BATCH_SIZE = 6
    processed = 0
    failed = 0

    for i in range(0, len(to_process), BATCH_SIZE):
        batch = to_process[i:i+BATCH_SIZE]
        tasks = [process_file(f, force=force) for f in batch]
        results = await asyncio.gather(*tasks)
        
        for success in results:
            if success:
                processed += 1
            else:
                failed += 1
        
        print(f"  Progress: {processed + failed}/{len(to_process)} ({failed} failed)\n")
        # Small delay between batches to respect rate limits
        await asyncio.sleep(1.0)

    print(f"\nExtraction complete: {processed} succeeded, {failed} failed.")
    print(f"Output directory: {OUTPUT_DIR}")

if __name__ == "__main__":
    asyncio.run(main())
