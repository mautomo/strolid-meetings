import os
import sys
import json
from pathlib import Path
from datetime import datetime

# Ensure we can import from src/
sys.path.insert(0, str(Path(__file__).resolve().parent))
from schema_types import Person, Decision, ActionItem, TopicThread, DirectionChange, PersonDirectionChange, ExtractedMeeting, NormalizedData

EXTRACTED_DIR = Path(__file__).resolve().parents[1] / "data" / "extracted"
OUTPUT_DIR = Path(__file__).resolve().parents[1] / "data" / "normalized"

NAME_ALIASES = {
    "vin": "Vinnie Micciche",
    "vinnie": "Vinnie Micciche",
    "vinnie m": "Vinnie Micciche",
    "vinnie micciche": "Vinnie Micciche",
    "michael": "Michael Donovan",
    "michael d": "Michael Donovan",
    "michael donovan": "Michael Donovan",
    "joe": "Joe Furnari",
    "joe f": "Joe Furnari",
    "joe furnari": "Joe Furnari",
    "jason": "Jason Branham",
    "jason branham": "Jason Branham",
    "matt": "Matt Watson",
    "matt watson": "Matt Watson",
    "paulo": "Paulo Trovao",
    "paulo trovao": "Paulo Trovao",
    "shawna": "Shawna Behen",
    "shawna behen": "Shawna Behen",
    "thomas": "Thomas Howe",
    "thomas howe": "Thomas Howe",
    "sergey": "Sergey",
    "sophia": "Sophia",
    "jake": "Jake",
    "link": "Link",
}

def resolve_name(name: str) -> str:
    key = name.strip().lower()
    return NAME_ALIASES.get(key, name.strip())

def make_person_id(name: str) -> str:
    # lower case alphanumeric with hyphens
    cleaned = "".join(c if c.isalnum() else "-" for c in name.lower())
    # replace duplicate hyphens
    while "--" in cleaned:
        cleaned = cleaned.replace("--", "-")
    return cleaned.strip("-")

def standardize_date(date_str: str) -> str:
    if not date_str:
        return "2025-01-01"
    date_str = date_str.strip()
    
    # Try %Y-%m-%d and other common patterns
    for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%B %d, %Y", "%b %d, %Y", "%d %B %Y", "%d %b %Y", "%m/%d/%Y", "%m-%d-%Y"]:
        try:
            dt = datetime.strptime(date_str, fmt)
            return dt.strftime("%Y-%m-%d")
        except ValueError:
            continue
            
    # Try parsing month names manually or using dateutil
    try:
        from dateutil import parser
        dt = parser.parse(date_str)
        return dt.strftime("%Y-%m-%d")
    except Exception:
        pass
        
    return date_str

import docx
import re
from pypdf import PdfReader

MEETINGS_DIR = Path(__file__).resolve().parents[2] / "original-docs"

def extract_text_from_file(stem: str) -> str:
    for ext in [".docx", ".md", ".pdf", ".txt"]:
        file_path = MEETINGS_DIR / f"{stem}{ext}"
        if file_path.exists():
            try:
                if ext == ".docx":
                    doc = docx.Document(file_path)
                    full_text = []
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    full_text.append(para.text)
                    return "\n".join(full_text)
                elif ext == ".pdf":
                    reader = PdfReader(file_path)
                    text_parts = []
                    for page in reader.pages:
                        t = page.extract_text()
                        if t:
                            text_parts.append(t)
                    return "\n".join(text_parts)
                else:
                    with open(file_path, "r", encoding="utf-8") as f:
                        return f.read()
            except Exception as e:
                print(f"Warning: Failed to read raw file {file_path}: {e}")
    return ""

def calculate_participation(meeting: ExtractedMeeting, stem: str):
    from schema_types import AttendeeParticipation
    
    text = extract_text_from_file(stem)
    if not text:
        return
        
    # Find transcript section
    transcript_start = -1
    for kw in ["Transcript", "transcript", "📖 Transcript", "📖 transcript"]:
        idx = text.find(kw)
        if idx != -1:
            transcript_start = idx
            break
            
    transcript_text = text[transcript_start:] if transcript_start != -1 else text
    
    # Analyze duration via timestamps (e.g. 00:48:19 or 48:19)
    timestamps = re.findall(r'(\d{1,2}):(\d{2}):(\d{2})', transcript_text)
    if not timestamps:
        timestamps = re.findall(r'(\d{1,2}):(\d{2})', transcript_text)
        
    duration = 0
    if timestamps:
        last_ts = timestamps[-1]
        if len(last_ts) == 3:
            h, m, s = map(int, last_ts)
            duration = h * 60 + m
        else:
            m, s = map(int, last_ts)
            duration = m
            
    meeting.durationMinutes = duration if duration > 0 else None

    # Count words spoken per attendee
    speaker_words = {resolve_name(a): 0 for a in meeting.attendees}
    
    # Build pattern matching for attendee names and first names/aliases
    name_patterns = {}
    for raw_name in meeting.attendees:
        full_name = resolve_name(raw_name)
        patterns = {full_name.lower(), raw_name.strip().lower()}
        parts = full_name.lower().split()
        if len(parts) > 0:
            patterns.add(parts[0]) 
        if len(parts) > 1:
            patterns.add(f"{parts[0]} {parts[1][0]}") 
        name_patterns[full_name] = patterns
        
    lines = transcript_text.splitlines()
    current_speaker = None
    
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue
            
        match = re.match(r'^\[?([A-Za-z\s._-]+)\]?:\s*(.*)', line_strip)
        if match:
            potential_speaker = match.group(1).strip().lower()
            content = match.group(2).strip()
            
            found = False
            for full_name, patterns in name_patterns.items():
                if potential_speaker in patterns:
                    current_speaker = full_name
                    found = True
                    break
            if not found:
                current_speaker = None
                
            if current_speaker:
                words = len(content.split())
                speaker_words[current_speaker] += words
        else:
            if current_speaker:
                words = len(line_strip.split())
                speaker_words[current_speaker] += words
                
    total_words = sum(speaker_words.values())
    
    participation_list = []
    for raw_name in meeting.attendees:
        full_name = resolve_name(raw_name)
        words = speaker_words.get(full_name, 0)
        pct = (words / total_words * 100) if total_words > 0 else 0.0
        
        level = "NONE"
        if pct > 25.0:
            level = "HIGH"
        elif pct > 10.0:
            level = "MEDIUM"
        elif words > 0:
            level = "LOW"
            
        participation_list.append(AttendeeParticipation(
            person_name=full_name,
            words_spoken=words,
            participation_percentage=pct,
            level=level
        ))
        
    meeting.participation = participation_list

def load_extracted_meetings() -> list[ExtractedMeeting]:
    if not EXTRACTED_DIR.exists():
        return []
    
    meetings = []
    for file in os.listdir(EXTRACTED_DIR):
        if not file.endswith(".json"):
            continue
        try:
            with open(EXTRACTED_DIR / file, "r", encoding="utf-8") as f:
                data = json.load(f)
                meeting = ExtractedMeeting(**data)
            
            # Standardize meeting date to YYYY-MM-DD
            meeting.date = standardize_date(meeting.date)
            
            # Calculate participation from original document text
            stem = Path(file).stem
            calculate_participation(meeting, stem)
            
            meetings.append(meeting)
        except Exception as e:
            print(f"Error loading {file}: {e}")
            
    # Sort chronologically by date
    meetings.sort(key=lambda m: m.date)
    return meetings


def build_people(meetings: list[ExtractedMeeting]) -> list[Person]:
    people_map = {}
    
    for meeting in meetings:
        for raw_name in meeting.attendees:
            name = resolve_name(raw_name)
            p_id = make_person_id(name)
            
            if p_id not in people_map:
                people_map[p_id] = Person(
                    id=p_id,
                    name=name,
                    aliases=[],
                    meetingCount=0
                )
                
            person = people_map[p_id]
            person.meetingCount += 1
            
            trimmed = raw_name.strip()
            if trimmed != name and trimmed not in person.aliases:
                person.aliases.append(trimmed)
                
    # Infer department/roles based on meeting attendance patterns
    for person in people_map.values():
        attended_types = set()
        for meeting in meetings:
            if any(resolve_name(a) == person.name for a in meeting.attendees):
                attended_types.add(meeting.type)
                
        if "leadership" in attended_types:
            person.department = "leadership"
        elif "marketing" in attended_types and "product" not in attended_types:
            person.department = "marketing"
        elif "product" in attended_types and "marketing" not in attended_types:
            person.department = "product"
        else:
            person.department = "cross-functional"
            
    # Return sorted by meeting count descending
    return sorted(people_map.values(), key=lambda p: p.meetingCount, reverse=True)

def build_decisions(meetings: list[ExtractedMeeting]) -> list[Decision]:
    decisions = []
    counter = 1
    
    for meeting in meetings:
        for d in meeting.decisions:
            decisions.append(Decision(
                id=f"d-{str(counter).zfill(3)}",
                description=d.description,
                decidedBy=[resolve_name(name) for name in d.decidedBy],
                meetingId=meeting.meetingId,
                meetingDate=meeting.date,
                topic=d.topic,
                confidence=d.confidence
            ))
            counter += 1
            
    # Chain decisions within the same topic
    by_topic = {}
    for d in decisions:
        by_topic.setdefault(d.topic, []).append(d)
        
    for topic, topic_decisions in by_topic.items():
        if len(topic_decisions) < 2:
            continue
        
        # Sort by date
        topic_decisions.sort(key=lambda x: x.meetingDate)
        
        # Mark lineage
        for i in range(1, len(topic_decisions)):
            prev = topic_decisions[i - 1]
            curr = topic_decisions[i]
            
            if prev.description.lower() != curr.description.lower():
                curr.supersedes = prev.id
                prev.supersededBy = curr.id
                
    return decisions

def build_action_items(meetings: list[ExtractedMeeting]) -> list[ActionItem]:
    items = []
    counter = 1
    
    for meeting in meetings:
        for ai in meeting.actionItems:
            expected = None
            if ai.expectedOutcome:
                expected = ai.expectedOutcome
            items.append(ActionItem(
                id=f"a-{str(counter).zfill(3)}",
                task=ai.task,
                owner=resolve_name(ai.owner),
                meetingId=meeting.meetingId,
                meetingDate=meeting.date,
                deadline=ai.deadline or None,
                status="open",
                expectedOutcome=expected
            ))
            counter += 1
            
    # Cross-reference completion in later meetings
    for item in items:
        # Get important words from task
        words = [w.lower() for w in item.task.split() if len(w) > 4]
        item_time = datetime.strptime(item.meetingDate, "%Y-%m-%d")
        
        for meeting in meetings:
            meeting_time = datetime.strptime(meeting.date, "%Y-%m-%d")
            if meeting_time <= item_time:
                continue
                
            # Scan meeting summaries, decisions, reference texts
            meeting_texts = [
                meeting.summary,
                *meeting.referencesToPast,
                *[d.description for d in meeting.decisions]
            ]
            full_meeting_text = " ".join(meeting_texts).lower()
            
            # Count word matches
            match_count = sum(1 for w in words if w in full_meeting_text)
            if match_count >= 3:
                # Check for completion markers
                completion_markers = ["completed", "done", "finished", "launched", "shipped"]
                if any(marker in full_meeting_text for marker in completion_markers):
                    item.status = "done"
                    item.resolvedInMeeting = meeting.meetingId
                    break
                    
        # Mark recurring/abandoned if still open and older than 60 days
        if item.status == "open":
            days_since = (datetime.now() - item_time).days
            if days_since > 60:
                # Check if it keeps appearing as a task in later meetings
                later_mentions = []
                for m in meetings:
                    m_time = datetime.strptime(m.date, "%Y-%m-%d")
                    if m_time <= item_time:
                        continue
                        
                    later_tasks_text = " ".join(a.task for a in m.actionItems).lower()
                    if sum(1 for w in words if w in later_tasks_text) >= 3:
                        later_mentions.append(m)
                        
                if len(later_mentions) >= 2:
                    item.status = "recurring"
                elif len(later_mentions) == 0:
                    item.status = "abandoned"
                    
    return items

def build_topics(meetings: list[ExtractedMeeting], decisions: list[Decision]) -> list[TopicThread]:
    topic_map = {}
    
    for meeting in meetings:
        for topic in meeting.topicsDiscussed:
            if topic not in topic_map:
                topic_map[topic] = TopicThread(
                    id=topic,
                    label=topic,
                    firstMentioned=meeting.date,
                    lastMentioned=meeting.date,
                    meetingRefs=[],
                    decisionRefs=[],
                    mentionCount=0
                )
                
            thread = topic_map[topic]
            thread.meetingRefs.append(meeting.meetingId)
            thread.lastMentioned = meeting.date
            thread.mentionCount += 1
            
    for d in decisions:
        if d.topic in topic_map:
            topic_map[d.topic].decisionRefs.append(d.id)
            
    # Return sorted by mention count descending
    return sorted(topic_map.values(), key=lambda t: t.mentionCount, reverse=True)

def detect_direction_changes(decisions: list[Decision]) -> list[DirectionChange]:
    changes = []
    
    by_topic = {}
    for d in decisions:
        by_topic.setdefault(d.topic, []).append(d)
        
    for topic, topic_decisions in by_topic.items():
        if len(topic_decisions) < 2:
            continue
            
        topic_decisions.sort(key=lambda x: x.meetingDate)
        
        for i in range(1, len(topic_decisions)):
            prev = topic_decisions[i - 1]
            curr = topic_decisions[i]
            
            if curr.supersedes == prev.id:
                prev_date = datetime.strptime(prev.meetingDate, "%Y-%m-%d")
                curr_date = datetime.strptime(curr.meetingDate, "%Y-%m-%d")
                days_between = (curr_date - prev_date).days
                
                changes.append(DirectionChange(
                    topic=topic,
                    originalPosition=prev.description,
                    originalDate=prev.meetingDate,
                    originalMeeting=prev.meetingId,
                    newPosition=curr.description,
                    changeDate=curr.meetingDate,
                    changeMeeting=curr.meetingId,
                    daysBetween=days_between
                ))
                
    return sorted(changes, key=lambda c: c.changeDate)

def detect_person_direction_changes(decisions: list[Decision]) -> list[PersonDirectionChange]:
    """Per-person reversals: the SAME individual sets a direction on a topic, then
    later supersedes it. We walk each topic's superseding chain (same logic as
    detect_direction_changes) and attribute the reversal to people present in BOTH
    the original decision's decidedBy and the superseding decision's decidedBy.
    """
    changes = []

    by_topic = {}
    for d in decisions:
        by_topic.setdefault(d.topic, []).append(d)

    for topic, topic_decisions in by_topic.items():
        if len(topic_decisions) < 2:
            continue

        topic_decisions.sort(key=lambda x: x.meetingDate)

        for i in range(1, len(topic_decisions)):
            prev = topic_decisions[i - 1]
            curr = topic_decisions[i]

            if curr.supersedes != prev.id:
                continue

            # Same individual on both sides of the reversal.
            shared = [name for name in curr.decidedBy if name in prev.decidedBy]
            if not shared:
                continue

            prev_date = datetime.strptime(prev.meetingDate, "%Y-%m-%d")
            curr_date = datetime.strptime(curr.meetingDate, "%Y-%m-%d")
            days_between = (curr_date - prev_date).days

            for person in shared:
                changes.append(PersonDirectionChange(
                    person=person,
                    topic=topic,
                    originalPosition=prev.description,
                    originalDate=prev.meetingDate,
                    originalMeeting=prev.meetingId,
                    originalConfidence=prev.confidence,
                    newPosition=curr.description,
                    changeDate=curr.meetingDate,
                    changeMeeting=curr.meetingId,
                    newConfidence=curr.confidence,
                    daysBetween=days_between
                ))

    return sorted(changes, key=lambda c: (c.person, c.changeDate))

def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    print("Loading extracted meetings...")
    meetings = load_extracted_meetings()
    print(f"Loaded {len(meetings)} meetings.\n")
    
    if not meetings:
        print("No meetings loaded. Exiting.")
        return
        
    print("Building people index...")
    people = build_people(meetings)
    print(f"  Found {len(people)} unique people.\n")
    
    print("Building decision log...")
    decisions = build_decisions(meetings)
    print(f"  Found {len(decisions)} decisions.\n")
    
    print("Building action item tracker...")
    action_items = build_action_items(meetings)
    status_counts = {}
    for ai in action_items:
        status_counts[ai.status] = status_counts.get(ai.status, 0) + 1
    print(f"  Found {len(action_items)} action items: {status_counts}\n")
    
    print("Building topic threads...")
    topics = build_topics(meetings, decisions)
    print(f"  Found {len(topics)} unique topics.\n")
    
    print("Detecting direction changes...")
    direction_changes = detect_direction_changes(decisions)
    print(f"  Found {len(direction_changes)} direction changes.\n")

    print("Detecting per-person direction reversals...")
    person_direction_changes = detect_person_direction_changes(decisions)
    print(f"  Found {len(person_direction_changes)} per-person reversals.\n")

    normalized = NormalizedData(
        people=people,
        decisions=decisions,
        actionItems=action_items,
        topics=topics,
        directionChanges=direction_changes,
        personDirectionChanges=person_direction_changes,
        meetings=meetings
    )
    
    # Save individual JSON files
    with open(OUTPUT_DIR / "people.json", "w", encoding="utf-8") as f:
        json.dump([p.model_dump() for p in people], f, indent=2)
    with open(OUTPUT_DIR / "decisions.json", "w", encoding="utf-8") as f:
        json.dump([d.model_dump() for d in decisions], f, indent=2)
    with open(OUTPUT_DIR / "actions.json", "w", encoding="utf-8") as f:
        json.dump([a.model_dump() for a in action_items], f, indent=2)
    with open(OUTPUT_DIR / "topics.json", "w", encoding="utf-8") as f:
        json.dump([t.model_dump() for t in topics], f, indent=2)
    with open(OUTPUT_DIR / "direction-changes.json", "w", encoding="utf-8") as f:
        json.dump([c.model_dump() for c in direction_changes], f, indent=2)
    with open(OUTPUT_DIR / "person-direction-changes.json", "w", encoding="utf-8") as f:
        json.dump([c.model_dump() for c in person_direction_changes], f, indent=2)
    with open(OUTPUT_DIR / "meetings.json", "w", encoding="utf-8") as f:
        json.dump([m.model_dump() for m in meetings], f, indent=2)
        
    # Save combined file
    with open(OUTPUT_DIR / "all.json", "w", encoding="utf-8") as f:
        f.write(normalized.model_dump_json(indent=2))
        
    print("Normalization complete.")
    print(f"Output directory: {OUTPUT_DIR}")

if __name__ == "__main__":
    main()
